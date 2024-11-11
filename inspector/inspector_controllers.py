from fastapi import APIRouter, HTTPException, Depends
from sqlalchemy.orm import Session
from typing import List
from inspector.inspector_models import Inspector
from inspector.inspector_schemas import InspectorCreate, InspectorUpdate, InspectorResponse
from database import get_db
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from report.report_models import Report

# Создаем HTTPBearer схему для проверки токена
oauth2_scheme = HTTPBearer()

# Инициализируем роутер
router = APIRouter()

# CRUD операции для инспекторов

# 1. Создание инспектора
@router.post("/", response_model=InspectorResponse, dependencies=[Depends(oauth2_scheme)])
def create_inspector(inspector: InspectorCreate, db: Session = Depends(get_db)):
    new_inspector = Inspector(**inspector.model_dump())  # Используем model_dump вместо dict
    db.add(new_inspector)
    db.commit()
    db.refresh(new_inspector)
    return new_inspector

# 2. Чтение всех инспекторов
@router.get("/", response_model=List[InspectorResponse], dependencies=[Depends(oauth2_scheme)])
def read_inspectors(skip: int = 0, limit: int = 10, db: Session = Depends(get_db)):
    inspectors = db.query(Inspector).offset(skip).limit(limit).all()
    return inspectors

# 3. Чтение инспектора по ID
@router.get("/{inspector_id}", response_model=InspectorResponse, dependencies=[Depends(oauth2_scheme)])
def read_inspector(inspector_id: int, db: Session = Depends(get_db)):
    inspector = db.query(Inspector).filter(Inspector.id == inspector_id).first()
    if inspector is None:
        raise HTTPException(status_code=404, detail="Inspector not found")
    return inspector

# 4. Обновление инспектора
@router.put("/{inspector_id}", response_model=InspectorResponse, dependencies=[Depends(oauth2_scheme)])
def update_inspector(inspector_id: int, inspector_update: InspectorUpdate, db: Session = Depends(get_db)):
    db_inspector = db.query(Inspector).filter(Inspector.id == inspector_id).first()
    if db_inspector is None:
        raise HTTPException(status_code=404, detail="Inspector not found")

    for key, value in inspector_update.model_dump(exclude_unset=True).items():
        setattr(db_inspector, key, value)

    db.commit()
    db.refresh(db_inspector)
    return db_inspector

# 5. Удаление инспектора
@router.delete("/{inspector_id}", dependencies=[Depends(oauth2_scheme)])
def delete_inspector(inspector_id: int, db: Session = Depends(get_db)):
    db_inspector = db.query(Inspector).filter(Inspector.id == inspector_id).first()
    if db_inspector is None:
        raise HTTPException(status_code=404, detail="Inspector not found")

    db.delete(db_inspector)
    db.commit()
    return {"detail": "Inspector deleted"}

# Функция для получения ФИО в дательном падеже
def get_full_name_dative(inspector):
    first_name = inspector.first_name
    last_name = inspector.last_name
    patronymic = inspector.patronymic

    if last_name.endswith('ова') or last_name.endswith('ева'):
        last_name_dative = last_name[:-1] + "е"
    elif last_name.endswith('ин') or last_name.endswith('ов') or last_name.endswith('ев'):
        last_name_dative = last_name + "у"
    else:
        last_name_dative = last_name + "у"

    if first_name.endswith('а'):
        first_name_dative = first_name
    else:
        first_name_dative = first_name + "у"

    if patronymic.endswith('на'):
        patronymic_dative = patronymic[:-1] + "не"
    elif patronymic.endswith('ич'):
        patronymic_dative = patronymic + "у"
    else:
        patronymic_dative = patronymic + "у"

    return f"{last_name_dative} {first_name_dative} {patronymic_dative}"

# Функция для получения должности в дательном падеже
def get_position_dative(position):
    if position == 'Инспектор':
        return "Инспектору"
    elif position == 'Главный инспектор':
        return "Главному инспектору"
    elif position == 'Старший инспектор':
        return "Старшему инспектору"
    return position

# Генерация документа Word с информацией об инспекторах и отчетах
@router.get("/report/", dependencies=[Depends(oauth2_scheme)])
def generate_inspector_report(month: int, year: int, db: Session = Depends(get_db)):
    if month < 1 or month > 12:
        raise HTTPException(status_code=400, detail="Некорректный месяц")

    reports = db.query(Report).filter(
        Report.report_date.like(f"{year}-{month:02}-%")
    ).all()

    if not reports:
        raise HTTPException(status_code=404, detail="Нет отчетов за указанный месяц и год")

    try:
        doc = Document("report_template.docx")
    except Exception:
        raise HTTPException(status_code=500, detail="Шаблон не найден или ошибка при загрузке шаблона")

    months_russian = [
        "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
        "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"
    ]

    for paragraph in doc.paragraphs:
        if "{{date}}" in paragraph.text:
            current_date = datetime.now().strftime("%d.%m.%Y")
            paragraph.text = paragraph.text.replace("{{date}}", current_date)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        if "{{month_year}}" in paragraph.text:
            month_year_str = f"{months_russian[month - 1]} {year} года"
            paragraph.text = paragraph.text.replace("{{month_year}}", month_year_str)
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

        if "{{inspectors}}" in paragraph.text:
            inspectors_info = ""
            for report in reports:
                inspector = db.query(Inspector).filter(Inspector.id == report.inspector_id).first()
                if inspector:
                    full_name = get_full_name_dative(inspector)
                    position = get_position_dative(inspector.position)
                    inspectors_info += f"{full_name} – {position}\n"

            paragraph.text = ""
            run = paragraph.add_run(inspectors_info)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=inspectors_report.docx"}
    )
